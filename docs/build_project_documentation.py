from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_FILE = Path(__file__).with_name("HappyMeals_Project_Documentation.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        set_cell_shading(hdr[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 14)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(4)


doc = Document()
configure_styles(doc)

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("Министерство науки и высшего образования Российской Федерации\n")
r.bold = True
r.font.name = "Times New Roman"
r.font.size = Pt(14)
p.add_run("Федеральное государственное образовательное учреждение высшего образования\n")
p.add_run("Кафедра информационных технологий\n")

doc.add_paragraph("\n\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ\n")
r.bold = True
r.font.size = Pt(18)
r.font.name = "Times New Roman"
p.add_run("к клиент-серверному приложению\n").font.size = Pt(14)
r = p.add_run("«Шеф-помощник» (HappyMeals)")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

doc.add_paragraph("\n\n")
meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
for left, right, row in [
    ("Выполнил", "______________________________", 0),
    ("Группа", "______________________________", 1),
    ("Проверил", "______________________________", 2),
    ("Город, год", "______________________________, 2026", 3),
]:
    set_cell_text(meta.rows[row].cells[0], left, bold=True)
    set_cell_text(meta.rows[row].cells[1], right)

doc.add_page_break()

# Contents
doc.add_heading("Содержание", level=1)
contents = [
    "Введение",
    "1 Общие сведения о проекте",
    "2 Анализ предметной области",
    "3 Требования к программному продукту",
    "4 Проектирование программного обеспечения",
    "5 Описание реализации",
    "6 Руководство пользователя",
    "7 Руководство программиста и администратора",
    "8 Контрольный пример и тестирование",
    "Заключение",
    "Список использованных источников",
    "Приложение А. Структура файлов проекта",
]
for item in contents:
    p = doc.add_paragraph(item)
    p.paragraph_format.first_line_indent = Cm(0)

doc.add_page_break()

doc.add_heading("Введение", level=1)
doc.add_paragraph(
    "Настоящий документ описывает клиент-серверное приложение «Шеф-помощник» "
    "(HappyMeals), предназначенное для подбора рецептов по пользовательским "
    "предпочтениям. Документация подготовлена в формате пояснительной записки: "
    "от постановки задачи и требований до описания архитектуры, базы данных, "
    "пользовательских сценариев, сборки, запуска и тестирования."
)
doc.add_paragraph(
    "Актуальность разработки связана с необходимостью быстро подбирать блюда "
    "по набору ограничений: исключаемым ингредиентам, кухне, типу блюда, времени "
    "приготовления и сложности. Разделение приложения на сервер и клиент позволяет "
    "разместить серверную часть на VDS и подключать к ней клиентов с разных устройств."
)

doc.add_heading("1 Общие сведения о проекте", level=1)
doc.add_heading("1.1 Наименование и назначение", level=2)
doc.add_paragraph(
    "Наименование программного продукта: «Шеф-помощник» (HappyMeals). Программа "
    "предназначена для регистрации пользователей, авторизации, подбора рецептов, "
    "просмотра подробного описания блюд, ведения избранного, истории поиска и "
    "пользовательской статистики."
)
doc.add_heading("1.2 Цели разработки", level=2)
add_bullets(
    doc,
    [
        "реализовать независимый запуск клиента и сервера на разных устройствах;",
        "разграничить пользовательский интерфейс и серверную бизнес-логику;",
        "сохранять пользовательские данные на сервере и привязывать их к аккаунту;",
        "обеспечить подбор рецептов по нескольким фильтрам;",
        "подготовить проект к размещению серверной части на VDS.",
    ],
)

doc.add_heading("1.3 Используемые технологии", level=2)
add_table(
    doc,
    ["Компонент", "Технология", "Назначение"],
    [
        ["Язык программирования", "C++17", "Основная реализация клиента и сервера"],
        ["Фреймворк", "Qt 6", "GUI, TCP-сеть, JSON, SQLite"],
        ["Сеть", "QTcpServer, QTcpSocket", "Обмен между клиентом и сервером"],
        ["База данных", "SQLite", "Хранение пользователей, рецептов, истории, избранного и статистики"],
        ["Система сборки", "qmake, MinGW", "Сборка отдельных .pro-проектов клиента и сервера"],
        ["ОС целевого запуска", "Windows Server / Windows", "Запуск сервера на VDS и клиента на ПК пользователя"],
    ],
    widths=[4, 4, 8],
)

doc.add_heading("2 Анализ предметной области", level=1)
doc.add_paragraph(
    "Предметная область приложения — подбор кулинарных рецептов. Пользователь "
    "может не помнить конкретное блюдо, но может задать ограничения: не использовать "
    "определенные ингредиенты, выбрать кухню, тип блюда, максимальное время приготовления "
    "и сложность. Система должна подобрать подходящие рецепты и показать их в удобном виде."
)
doc.add_heading("2.1 Основные участники", level=2)
add_table(
    doc,
    ["Участник", "Описание", "Основные действия"],
    [
        ["Гость", "Пользователь до авторизации", "Регистрация, переход к форме входа"],
        ["Пользователь", "Авторизованный клиент приложения", "Поиск рецептов, просмотр деталей, работа с избранным, история и статистика"],
        ["Сервер", "Центральная часть системы", "Проверка учетных данных, фильтрация рецептов, хранение пользовательских данных"],
        ["Администратор VDS", "Ответственный за размещение серверной части", "Сборка, запуск, настройка порта и firewall"],
    ],
    widths=[3.2, 5.4, 7.2],
)

doc.add_heading("3 Требования к программному продукту", level=1)
doc.add_heading("3.1 Функциональные требования", level=2)
add_table(
    doc,
    ["№", "Требование", "Описание"],
    [
        ["FR-01", "Регистрация", "Система должна создавать пользователя с уникальным логином, паролем и email."],
        ["FR-02", "Авторизация", "Система должна разрешать вход только по существующему логину и корректному паролю."],
        ["FR-03", "Валидация форм", "Клиент должен проверять формат логина, пароля и email до отправки запроса на сервер."],
        ["FR-04", "Проверка дубля логина", "Сервер должен отклонять регистрацию, если логин уже существует в базе."],
        ["FR-05", "Подбор рецептов", "Система должна искать блюда по ингредиентам-исключениям, кухне, типу, времени и сложности."],
        ["FR-06", "Подробности рецепта", "Пользователь должен видеть название, кухню, время и описание приготовления."],
        ["FR-07", "Избранное", "Пользователь должен добавлять и удалять блюда из избранного."],
        ["FR-08", "История поиска", "История должна быть читаемой, без подряд идущих дублей и привязанной к аккаунту."],
        ["FR-09", "Статистика", "Система должна показывать количество поисков, избранных рецептов и время в приложении."],
        ["FR-10", "Выход из аккаунта", "При выходе клиент очищает локальное состояние, а сервер завершает пользовательскую сессию."],
    ],
    widths=[2, 4.2, 9.8],
)

doc.add_heading("3.2 Нефункциональные требования", level=2)
add_bullets(
    doc,
    [
        "клиент и сервер должны запускаться независимо;",
        "сервер должен принимать подключения по TCP на всех IPv4-интерфейсах;",
        "пользовательские данные не должны зависеть от конкретного устройства или текущей сессии;",
        "сетевой обмен должен выполняться в структурированном JSON-формате;",
        "данные должны сохраняться между запусками программы в SQLite;",
        "интерфейс клиента должен отображать блюда отдельными элементами, а не сплошным текстом.",
    ],
)

doc.add_heading("3.3 Входные и выходные данные", level=2)
add_table(
    doc,
    ["Тип данных", "Пример", "Назначение"],
    [
        ["Входные данные регистрации", "login, password, email", "Создание учетной записи"],
        ["Входные данные авторизации", "login, password", "Получение userId и начало сессии"],
        ["Параметры поиска", "excludedIngredients, cuisines, dishTypes, maxTime, maxComplexity", "Формирование запроса подбора рецептов"],
        ["Выходные данные поиска", "name, prepTime", "Список найденных блюд"],
        ["Подробности блюда", "html", "Форматированное описание рецепта"],
        ["Статистика", "searches, favorites, app time", "Информация об активности аккаунта"],
    ],
    widths=[4.2, 5.2, 6.2],
)

doc.add_heading("4 Проектирование программного обеспечения", level=1)
doc.add_heading("4.1 Архитектура", level=2)
doc.add_paragraph(
    "Проект реализован как две самостоятельные программы: HappyMealsClient и "
    "HappyMealsServer. Клиент содержит графический интерфейс и отправляет серверу "
    "запросы. Сервер содержит бизнес-логику, работает с базами SQLite и возвращает "
    "результаты клиенту. Такое разделение позволяет размещать сервер на VDS, а клиент "
    "запускать на пользовательских компьютерах."
)
add_code_block(
    doc,
    "HappyMealsClient  --TCP/JSON-->  HappyMealsServer  --SQLite-->  HappyMealsDB.sqlite\n"
    "        |                               |\n"
    "        |                               +--> dishes.sqlite\n"
    "        +--> Qt Widgets GUI",
)

doc.add_heading("4.2 Модули системы", level=2)
add_table(
    doc,
    ["Модуль", "Расположение", "Ответственность"],
    [
        ["ClientApi", "src/clientapi.cpp, include/clientapi.h", "TCP-подключение, JSON-запросы, хранение текущего userId"],
        ["ClientSessionManager", "include/clientsessionmanager.h", "Локальное состояние клиента: текущий пользователь, избранное, история"],
        ["MainWindow", "src/mainwindow.cpp", "Формы входа, регистрации, поиска, результатов, статистики"],
        ["DishServer", "src/dishserver.cpp", "TCP-сервер, прием подключений и передача запросов в обработчик"],
        ["FunctionsToServer", "src/functionstoserver.cpp", "Разбор JSON/строковых команд и маршрутизация по менеджерам"],
        ["UserManager", "src/usermanager.cpp", "Регистрация, авторизация, избранное, история, статистика, logout"],
        ["DishManager", "src/dishmanager.cpp", "Подбор блюд и получение подробного описания рецепта"],
        ["DB_Singleton", "src/db_singleton.cpp", "Подключение к SQLite, таблицы, запросы, фильтрация рецептов"],
    ],
    widths=[3.6, 5.3, 7.1],
)

doc.add_heading("4.3 Структура базы данных", level=2)
add_table(
    doc,
    ["База", "Таблица", "Назначение"],
    [
        ["HappyMealsDB.sqlite", "users", "Учетные записи пользователей: id, login, password, email"],
        ["HappyMealsDB.sqlite", "history", "История поиска по user_id: текст запроса и дата создания"],
        ["HappyMealsDB.sqlite", "favorites", "Избранные рецепты пользователя, уникальные по паре user_id + dish_name"],
        ["HappyMealsDB.sqlite", "user_sessions", "Сессии пользователя и длительность работы в приложении"],
        ["dishes.sqlite", "dishes", "Справочник блюд: название, кухня, тип, время, сложность, инструкция"],
        ["dishes.sqlite", "ingredients", "Справочник ингредиентов"],
        ["dishes.sqlite", "dish_ingredients", "Связь многие-ко-многим между блюдами и ингредиентами"],
    ],
    widths=[4.1, 3.5, 8.2],
)
doc.add_paragraph(
    "Логика socket_id из пользовательских данных не используется: история, избранное "
    "и статистика привязываются к постоянному идентификатору users.id. Это необходимо, "
    "чтобы пользователь видел свои данные при входе с разных устройств."
)

doc.add_heading("4.4 Протокол взаимодействия", level=2)
doc.add_paragraph(
    "Основной формат обмена — UTF-8 JSON. После успешной авторизации сервер возвращает "
    "userId, и дальнейшие команды клиента автоматически отправляются с этим идентификатором."
)
add_table(
    doc,
    ["Команда", "Основные поля", "Результат"],
    [
        ["auth", "login, password", "ok, message, userId"],
        ["reg", "login, password, email", "ok, message"],
        ["get_dish", "userId, excludedIngredients, cuisines, dishTypes, maxTime, maxComplexity, summary", "массив dishes"],
        ["dish_details", "name", "html с подробным описанием"],
        ["get_stat", "userId", "строка статистики"],
        ["get_history", "userId", "массив history"],
        ["get_favorites", "userId", "массив favorites"],
        ["add_favorite", "userId, name", "ok, message"],
        ["remove_favorite", "userId, name", "ok, message"],
        ["logout", "userId", "завершение сессии"],
    ],
    widths=[3.2, 7.5, 5.1],
)

doc.add_heading("5 Описание реализации", level=1)
doc.add_heading("5.1 Клиентская часть", level=2)
doc.add_paragraph(
    "Клиентская часть реализована на Qt Widgets. Она отвечает за формы входа и "
    "регистрации, экран фильтров, список результатов, карточку рецепта, избранное, "
    "историю и статистику. Клиент не работает напрямую с базой данных: все операции, "
    "связанные с пользователями и рецептами, отправляются серверу."
)
doc.add_heading("5.2 Серверная часть", level=2)
doc.add_paragraph(
    "Серверная часть запускается как консольное приложение. Сервер слушает TCP-порт "
    "40000 на всех IPv4-интерфейсах, принимает запросы клиентов, разбирает команды, "
    "выполняет операции с SQLite и возвращает JSON-ответы."
)
doc.add_heading("5.3 Алгоритм подбора рецептов", level=2)
add_numbered(
    doc,
    [
        "Клиент формирует набор фильтров: исключаемые ингредиенты, кухня, тип блюда, время и сложность.",
        "Сервер сохраняет читаемое описание поискового запроса в историю пользователя.",
        "DB_Singleton выбирает блюда из dishes.sqlite с учетом заданных ограничений.",
        "Результаты возвращаются клиенту как отдельные элементы с названием и временем приготовления.",
        "При выборе блюда клиент запрашивает подробное описание рецепта.",
    ],
)

doc.add_heading("6 Руководство пользователя", level=1)
doc.add_heading("6.1 Запуск клиента", level=2)
doc.add_paragraph(
    "Клиент можно запускать с параметрами IP-адреса и порта сервера. Если параметры "
    "не указаны, используются значения по умолчанию, заданные в коде."
)
add_code_block(doc, "HappyMealsClient.exe 195.2.81.149 40000")
doc.add_heading("6.2 Основные сценарии", level=2)
add_numbered(
    doc,
    [
        "Открыть клиентское приложение.",
        "Зарегистрироваться или войти под существующим аккаунтом.",
        "Перейти к фильтрам рецептов.",
        "Выбрать исключаемые ингредиенты, кухню, тип блюда, максимальное время и сложность.",
        "Нажать кнопку поиска и просмотреть список найденных рецептов.",
        "Открыть рецепт, изучить подробное описание и при необходимости добавить его в избранное.",
        "Открыть историю поиска или статистику пользователя.",
        "Выйти из аккаунта.",
    ],
)

doc.add_heading("7 Руководство программиста и администратора", level=1)
doc.add_heading("7.1 Сборка сервера", level=2)
add_code_block(
    doc,
    "cd C:\\Users\\Administrator\\Desktop\\timp_srv\n"
    "set PATH=C:\\Qt\\6.11.1\\mingw_64\\bin;C:\\Qt\\Tools\\mingw1310_64\\bin;%PATH%\n"
    "qmake HappyMealsServer.pro\n"
    "mingw32-make\n"
    "windeployqt release\\HappyMealsServer.exe",
)
doc.add_heading("7.2 Запуск сервера", level=2)
add_code_block(doc, "release\\HappyMealsServer.exe")
doc.add_paragraph(
    "После запуска в консоли должно появиться сообщение о старте сервера на порту 40000. "
    "На VDS необходимо разрешить входящие TCP-подключения на этот порт и убедиться, "
    "что VPN или системный прокси не перехватывает локальные сетевые подключения."
)
doc.add_heading("7.3 Файлы для размещения сервера", level=2)
add_bullets(
    doc,
    [
        "release/HappyMealsServer.exe и DLL-файлы, добавленные windeployqt;",
        "папка data с HappyMealsDB.sqlite и dishes.sqlite;",
        "папки Qt-плагинов: sqldrivers, networkinformation, tls;",
        "run_server.bat для удобного запуска;",
        "при необходимости исходники из server_package для пересборки на сервере.",
    ],
)

doc.add_heading("8 Контрольный пример и тестирование", level=1)
doc.add_paragraph(
    "Для проверки работы приложения подготовлен QA-файл с тест-планом, чек-листом, "
    "тест-кейсом, дефектом и отчетом: qa/HappyMeals_QA_TestCase_Defect_v2.xlsx. "
    "Чек-лист содержит 36 проверок по регистрации, авторизации, соединению с сервером, "
    "фильтрам, результатам поиска, избранному, истории и статистике."
)
doc.add_heading("8.1 Контрольный пример", level=2)
add_table(
    doc,
    ["Шаг", "Действие", "Ожидаемый результат"],
    [
        ["1", "Запустить сервер HappyMealsServer", "Сервер сообщает о запуске на порту 40000"],
        ["2", "Запустить клиент с IP и портом сервера", "Открывается форма входа"],
        ["3", "Зарегистрировать пользователя", "Пользователь создан, после регистрации отображается форма входа"],
        ["4", "Войти под созданным пользователем", "Открывается главное меню приложения"],
        ["5", "Выполнить поиск с кухней RUSSIAN и временем до 60 минут", "Отображается список подходящих рецептов"],
        ["6", "Открыть рецепт", "Появляется подробное описание блюда"],
        ["7", "Добавить рецепт в избранное", "Рецепт сохраняется за текущим аккаунтом"],
        ["8", "Открыть историю и статистику", "История и статистика соответствуют текущему пользователю"],
    ],
    widths=[1.5, 6.4, 7.8],
)
doc.add_heading("8.2 Известный дефект", level=2)
doc.add_paragraph(
    "В QA-документации зафиксирован дефект High: история поиска может отображаться "
    "в техническом виде и дублироваться после повторного одинакового запроса. До "
    "подтвержденной повторной проверки дефект должен оставаться в статусе Opened/Retest, "
    "а не Fixed."
)

doc.add_heading("Заключение", level=1)
doc.add_paragraph(
    "В рамках проекта выполнено разделение приложения на клиентскую и серверную части. "
    "Серверная часть содержит бизнес-логику, работу с базой данных и сетевой протокол, "
    "а клиентская часть отвечает за пользовательский интерфейс. Пользовательские данные "
    "привязаны к user_id, что позволяет сохранять историю, избранное и статистику на "
    "сервере независимо от устройства пользователя."
)
doc.add_paragraph(
    "Дальнейшее развитие проекта может включать хеширование паролей, настройку HTTPS/TLS "
    "или защищенного канала, расширение набора рецептов, добавление ролей администратора, "
    "улучшение поиска по ингредиентам и автоматизированные интеграционные тесты клиента и сервера."
)

doc.add_heading("Список использованных источников", level=1)
sources = [
    "Шаблон и пример структуры проектной документации: https://studfile.net/preview/15933472/",
    "Документация Qt: https://doc.qt.io/",
    "SQLite Documentation: https://www.sqlite.org/docs.html",
    "Материалы проекта HappyMeals: исходные файлы src, include, .pro и QA-документы.",
]
for source in sources:
    p = doc.add_paragraph(style="List Number")
    p.add_run(source)

doc.add_heading("Приложение А. Структура файлов проекта", level=1)
add_code_block(
    doc,
    "HappyMealsClient.pro      проект клиентского приложения\n"
    "HappyMealsServer.pro      проект серверного приложения\n"
    "src/                      исходные файлы C++\n"
    "include/                  заголовочные файлы\n"
    "data/                     SQLite-базы данных\n"
    "server_package/           серверные файлы для копирования на VDS\n"
    "qa/                       тест-кейсы, дефекты и отчет по тестированию\n"
    "run_client.bat            запуск клиента\n"
    "run_server.bat            запуск сервера",
)

for section in doc.sections:
    add_page_number(section)

doc.save(OUT_FILE)
print(OUT_FILE.resolve())
